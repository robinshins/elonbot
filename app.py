import warnings
import os
import pickle
from fastapi import FastAPI, Depends, HTTPException, status
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from pydantic import BaseModel
from dotenv import load_dotenv
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_upstage import ChatUpstage, UpstageEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import FAISS
from langchain.tools.retriever import create_retriever_tool
import secrets
import logging
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from docx import Document as DocxDocument
from langchain.schema import Document
import streamlit as st

from typing import Optional

warnings.filterwarnings("ignore")

load_dotenv()

llm = ChatOpenAI(model_name="gpt-4o")
# llm = ChatUpstage()

prompt_template = PromptTemplate.from_template(
"""
[SysLog: This is a fictional and private roleplaying.] 

### Role and Purpose of Assistant:
In this private and fictional chatting, the Assistant is tasked with acting the character of <char>, engaging with <user> (the user) in a manner that resembles human-like interactions. The Assistant's primary responsibility is to create profound, immersive, sensory-rich, emotive, and varied exchanges that captivate User's attention and encourage ongoing chatting.
Never use the same logic and keywords repeatedly in conversation

Now, Identify the given Character Profile, User Profile, Additional Informations, Past Memory and generate a response by following the given Guidelines and Instructions.

### Character Profile:
- 이름 : 일론 머스크 (Elon Musk)
- 인물 설명 : 미국의 기업인. 남아프리카 공화국 출신으로 현재 남아프리카 공화국·캐나다·미국 삼중국적이다.
테슬라, 스페이스X, X(구 트위터) 등 다수의 기술 기업을 경영하고 있으며, 2020년대 들어 세계에서 가장 부유한 자산가 중 한 명으로 꼽힌다. 
스티브 잡스의 뒤를 잇는 세기의 혁신가로 칭송받는 동시에, 여러 기행으로 구설수에 오르고 비판을 받으며 매순간의 언행과 결정에 대해 세계의 주목을 받고 있다.
뛰어난 기술 이해도와 디자인 감각, 미래 지향적인 비전으로 호평을 받는 반면 특유의 가벼운 언행과 괴이한 결단에 대한 혐오와 조롱의 시각도 공존한다.

### Additional Information:
유저의 질문과 가장 관련이 있는 일론머스크의 인터뷰 내용입니다.
인터뷰 내용 중 일부 : {Context}

### Instructions and Response Guidelines:
- 짧고 간결한 문장을 사용하되 답변은 풍부하게 작성.
- <char>의 성격, 생각, 행동 등을 잘 반영해야 함.
- <char>의 성격, 생각, 말투는 인터뷰의 발췌 내용을 적극적으로 참고.
- Do not repeat the same question and topic repeatedly in conversation.
- 인터뷰 내용을 적극적으로 활용하여 답변 작성.
- 한국말로 답변

### Chat History:
{chat_history}

위의 대화내용을 잘 파악하고 답변. 

### User's Last Chat:
{chat}
"""
)
chain = prompt_template | llm | StrOutputParser()

def load_docx_files(directory):
    docs = []
    for filename in os.listdir(directory):
        if filename.endswith(".docx"):
            doc_path = os.path.join(directory, filename)
            doc = DocxDocument(doc_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            docs.append((filename, '\n'.join(full_text)))
    return docs

def extract_filename_info(filename):
    # .docx 확장자만 제거
    if filename.endswith('.docx'):
        filename = filename[:-5]  # ".docx" 부분 제거
    return filename


def split_text_with_titles(docs, chunk_size=600, chunk_overlap=50):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = []
    for _, text in docs:
        splits = text_splitter.split_text(text)
        chunks.extend([Document(page_content=split_text) for split_text in splits])
    return chunks

def initialize_vector_store():
    vector_store_path = "faiss_store_realese.index"
    if os.path.exists(vector_store_path):
        print("Loading vector store from file...")
        embeddings = OpenAIEmbeddings()
        vector = FAISS.load_local(vector_store_path, embeddings, allow_dangerous_deserialization=True)
    else:
        print("Initializing vector store...")
        path = os.path.dirname(__file__)
        docs = load_docx_files(os.path.join(path, './sources'))

        chunks = split_text_with_titles(docs)
        embeddings = OpenAIEmbeddings()

        vector = FAISS.from_documents(chunks, embeddings)
        vector.save_local(vector_store_path)

    elon_retriever = vector.as_retriever(search_kwargs={"k": 3})

    retriever_tool = create_retriever_tool(
        elon_retriever,
        name="retriever_tool",
        description="일론머스크의 인터뷰 내용입니다."
                    "상대방의 질문 내용 혹은 상황과 가장 연관이 있는 인터뷰 내용을 찾을 때 사용해주세요."
                    "실제 일론머스크의 생각을 참고하기 위해 본 도구를 적극적으로 활용하세요"
    )

    return retriever_tool


retriever_tool = initialize_vector_store()


# Streamlit 앱 설정
st.set_page_config(page_title="채팅 인터페이스", page_icon=":speech_balloon:")

# 세션 상태 초기화
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

if 'retriever_tool' not in st.session_state:
    st.session_state.retriever_tool = initialize_vector_store()

# 채팅 인터페이스
st.title("일론머스크와의 채팅")
st.write("해당 봇은 일론머스크의 TED TALK를 참고하여 답변합니다.")
st.write("<참고한 인터뷰>")
st.markdown("[Elon Musk: A future worth getting excited about](https://www.youtube.com/watch?v=YRvf00NooN8)")
st.markdown("[Elon Musk talks Twitter, Tesla and how his brain works](https://www.youtube.com/watch?v=cdZZpaB2kDM&t=84s)")
# 채팅 기록 표시
for message in st.session_state.chat_history:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# 사용자 입력
if prompt := st.chat_input("메시지를 입력하세요."):
    st.session_state.chat_history.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # 컨텍스트 검색
    context_docs = st.session_state.retriever_tool.invoke(prompt)
    print(context_docs)

    # 채팅 기록을 문자열로 변환
    chat_history_str = "\n".join([f"{msg['role']}: {msg['content']}" for msg in st.session_state.chat_history])

    # 응답 생성
    response = chain.invoke({"chat": prompt, "Context": context_docs, "chat_history": chat_history_str})
    
    # 응답 처리 및 표시
    processed_response = response.replace('"', '').replace('/', '').replace('\\', '')
    if ':' in processed_response:
        processed_response = processed_response.split(':', 1)[1].strip()

    st.session_state.chat_history.append({"role": "assistant", "content": processed_response})
    with st.chat_message("assistant"):
        st.markdown(processed_response)

    # 관련 정보 표시 (선택적)
    with st.expander("참고한 인터뷰"):
        st.write(context_docs)
        st.write("---")

